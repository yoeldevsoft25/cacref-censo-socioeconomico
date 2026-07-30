from __future__ import annotations

import sqlite3
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DB_PATH = ROOT / "census.db"
LOGO_CACREF = ROOT / "public" / "Logo_Original.png"
LOGO_FUTPV = ROOT / "public" / "logo2.jpeg"
OUT_DOCX = DOCS / "CACREF_Salud_Propuesta_Directivos.docx"

NAVY = RGBColor(15, 35, 64)
BLUE = RGBColor(46, 116, 181)
RED = RGBColor(190, 18, 60)
GRAY = RGBColor(91, 104, 124)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
LIGHT_RED = "FDEEEF"
WHITE = "FFFFFF"


def get_metrics() -> dict:
    fallback = {
        "total": 0,
        "medicamentos": 0,
        "cirugias": 0,
        "familiares": 0,
        "alto": 0,
        "medio": 0,
        "bajo": 0,
        "score_promedio": 0,
        "calidad_promedio": 0,
        "aporte_promedio": 0,
    }
    if not DB_PATH.exists():
        return fallback
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        row = conn.execute(
            """
            SELECT
              COUNT(*) total,
              SUM(CASE WHEN requiere_medicamento_cronico = 1 THEN 1 ELSE 0 END) medicamentos,
              SUM(CASE WHEN requiere_cirugia = 1 THEN 1 ELSE 0 END) cirugias,
              SUM(CASE WHEN familiar_requiere_asistencia = 1 THEN 1 ELSE 0 END) familiares,
              SUM(CASE WHEN risk_level = 'ALTO' THEN 1 ELSE 0 END) alto,
              SUM(CASE WHEN risk_level = 'MEDIO' THEN 1 ELSE 0 END) medio,
              SUM(CASE WHEN risk_level = 'BAJO' THEN 1 ELSE 0 END) bajo,
              ROUND(AVG(score), 2) score_promedio,
              ROUND(AVG(calidad_vida_escala), 2) calidad_promedio,
              ROUND(AVG(capacidad_cuota), 2) aporte_promedio
            FROM census_submissions
            """
        ).fetchone()
    result = dict(fallback)
    result.update({k: row[k] if row[k] is not None else 0 for k in row.keys()})
    return result


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="D7DEE8", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_cell(cell, fill: str | None = None, bold=False, color=RGBColor(0, 0, 0), size=9.5):
    if fill:
        set_cell_shading(cell, fill)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold


def add_run(paragraph, text, size=11, bold=False, color=RGBColor(0, 0, 0), italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def add_paragraph(doc, text="", size=11, bold=False, color=RGBColor(0, 0, 0), after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    add_run(p, text, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        run = add_run(p, text, size=16, bold=True, color=BLUE)
    elif level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        run = add_run(p, text, size=13, bold=True, color=BLUE)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = add_run(p, text, size=12, bold=True, color=NAVY)
    return run


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_run(p, text, size=10.5, color=RGBColor(22, 31, 46))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_run(p, text, size=10.5, color=RGBColor(22, 31, 46))
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_table_width(table, [6.35])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color="C7D2E2")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, size=11, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    add_run(p2, body, size=10.2, color=RGBColor(31, 41, 55))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    set_table_width(table, [6.35 / len(metrics)] * len(metrics))
    for idx, (label, value, fill) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, fill)
        set_cell_borders(cell, color="D8DEE9")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        add_run(p, str(value), size=18, bold=True, color=NAVY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        add_run(p2, label, size=8.2, bold=True, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_matrix(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        style_cell(cell, fill=header_fill, bold=True, color=NAVY, size=9.3)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            style_cell(cells[idx], size=9.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header_p, "CACREF Salud | Propuesta ejecutiva", size=8.5, color=GRAY)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_p, "Documento de trabajo para directivos CACREF - Uso interno", size=8.5, color=GRAY)


def add_cover(doc: Document, metrics: dict):
    logo_table = doc.add_table(rows=1, cols=2)
    set_table_width(logo_table, [3.15, 3.15])
    for cell in logo_table.rows[0].cells:
        set_cell_borders(cell, color=WHITE, size="0")
    left, right = logo_table.rows[0].cells
    if LOGO_CACREF.exists():
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(LOGO_CACREF), width=Inches(0.82))
    if LOGO_FUTPV.exists():
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture(str(LOGO_FUTPV), width=Inches(0.82))

    add_paragraph(doc, "", after=24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "CACREF Salud", size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    add_run(p, "Censo socioeconómico, salud y bienestar para colaboradores PDVSA/FUTPV", size=14, color=GRAY)

    add_callout(
        doc,
        "Mensaje central",
        "La plataforma permite pasar de una ayuda reactiva y dispersa a una gestión ordenada, auditable y financieramente responsable: ayudar mejor, priorizar con evidencia y construir programas recurrentes sin comprometer el equilibrio económico de CACREF.",
        fill=LIGHT_BLUE,
    )

    add_metric_strip(
        doc,
        [
            ("Registros demo", metrics["total"], LIGHT_GRAY),
            ("Medicamentos", metrics["medicamentos"], LIGHT_BLUE),
            ("Cirugías", metrics["cirugias"], LIGHT_RED),
            ("Riesgo alto", metrics["alto"], LIGHT_RED),
        ]
    )

    add_paragraph(
        doc,
        "Preparado para Junta Directiva / Directivos CACREF",
        size=10.5,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=2,
    )
    add_paragraph(
        doc,
        f"Fecha: {date.today().strftime('%d/%m/%Y')} | Versión ejecutiva",
        size=9.5,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=20,
    )

    add_heading(doc, "Decisión solicitada", level=2)
    for item in [
        "Aprobar el uso del censo como herramienta institucional de diagnóstico y priorización.",
        "Autorizar un piloto controlado con responsables, reglas de acceso y tablero directivo.",
        "Ratificar aporte base de 2% para inscritos y definir topes, reserva y cálculos posteriores.",
        "Iniciar negociación con proveedores de medicamentos, laboratorios, clínicas y servicios asociados.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()


def add_pages(doc: Document, metrics: dict):
    add_heading(doc, "1. Objetivo y norte institucional")
    add_callout(
        doc,
        "Objetivo",
        "Crear una base única y confiable para identificar necesidades reales de trabajadores, afiliados CACREF y familiares, priorizar casos por urgencia administrativa y construir programas sostenibles de apoyo.",
        fill=LIGHT_BLUE,
    )
    add_paragraph(
        doc,
        "El norte no es entregar ayudas sin control. El norte es combinar solidaridad cooperativa con data, aporte base de 2% para inscritos, convenios y trazabilidad. Los préstamos u otros programas se calculan luego según política, capacidad y aprobación.",
    )
    add_matrix(
        doc,
        ["Situación actual", "Con CACREF Salud"],
        [
            ["Solicitudes dispersas por llamadas, planillas o mensajes", "Formulario único con datos comparables"],
            ["Decisiones caso a caso con información incompleta", "Scoring administrativo y cola priorizada"],
            ["Poca trazabilidad de revisión y aprobación", "Workflow, auditoría e historial por caso"],
            ["Riesgo de prometer más de lo financiable", "Topes, reserva técnica y decisión por comité"],
        ],
        [3.05, 3.3],
    )

    add_heading(doc, "2. Capacidad ya desplegada")
    for item in [
        "Formulario público de cuatro pasos: identidad, vinculación, economía, salud y calidad de vida.",
        "Dashboard administrativo con KPIs, filtros, detalle de casos, documentos, comentarios y workflow.",
        "Roles diferenciados: capturista, vocal, presidente y director.",
        "Consulta pública por cédula, transparencia agregada y políticas de privacidad.",
        "Backend local y producción: SQLite/Turso, Express/Netlify Functions, autenticación y bitácora.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Alcance correcto",
        "La plataforma no diagnostica, no sustituye evaluación médica y no aprueba beneficios automáticamente. Ordena información y prepara decisiones para el comité.",
        fill=LIGHT_RED,
        title_color=RED,
    )

    doc.add_page_break()

    add_heading(doc, "3. Datos, score y lógica de priorización")
    add_paragraph(
        doc,
        "El score es administrativo. Su función es ordenar la cola y transparentar por qué un caso requiere atención antes que otro. La decisión final sigue siendo humana, institucional y documentada.",
    )
    add_matrix(
        doc,
        ["Componente", "Uso en la decisión", "Resultado esperado"],
        [
            ["Antigüedad y vinculación", "Reconoce trayectoria y pertenencia cooperativa", "Prioridad con sentido institucional"],
            ["Capacidad de aporte", "Evita cuotas impagables", "Sostenibilidad y menor mora"],
            ["Medicamento crónico", "Identifica necesidad recurrente", "Programas de abastecimiento y convenios"],
            ["Cirugía/procedimiento", "Detecta casos de alto impacto", "Revisión por comité y topes"],
            ["Calidad de vida", "Mide vulnerabilidad percibida", "Semáforo social de bienestar"],
        ],
        [1.6, 2.55, 2.2],
    )
    add_paragraph(doc, "Salidas del sistema: riesgo bajo/medio/alto, recomendación administrativa, prioridad numérica, cuota máxima sugerida y estado del workflow.", size=10.5)

    add_heading(doc, "4. Lectura de la base de demostración")
    add_metric_strip(
        doc,
        [
            ("Total casos", metrics["total"], LIGHT_GRAY),
            ("Riesgo alto", metrics["alto"], LIGHT_RED),
            ("Riesgo medio", metrics["medio"], LIGHT_BLUE),
            ("Riesgo bajo", metrics["bajo"], LIGHT_GRAY),
        ]
    )
    add_matrix(
        doc,
        ["Indicador demo", "Valor", "Lectura directiva"],
        [
            ["Medicamentos crónicos", metrics["medicamentos"], "Demanda recurrente que puede negociarse por volumen"],
            ["Cirugías/procedimientos", metrics["cirugias"], "Casos de alto impacto que requieren topes y comité"],
            ["Familiares con asistencia", metrics["familiares"], "Carga familiar que afecta calidad de vida y estabilidad"],
            ["Calidad de vida promedio", metrics["calidad_promedio"], "Línea base para medir mejora trimestral"],
            ["Aporte mensual promedio", metrics["aporte_promedio"], "Referencia para calibrar programas sostenibles"],
        ],
        [2.0, 1.0, 3.35],
    )
    add_paragraph(doc, "Nota: los datos son de demostración local para validar capacidad operativa; la campaña real debe iniciar con piloto controlado.", size=9.5, color=GRAY)

    doc.add_page_break()

    add_heading(doc, "5. Modelo recurrente sin números rojos")
    add_callout(
        doc,
        "Regla financiera de control",
        "El 2% opera como aporte base de inscritos. Ingresos recurrentes - ayudas directas - costos operativos - reserva técnica >= 0. Préstamos y programas especiales se calculan aparte.",
        fill=LIGHT_BLUE,
    )
    add_matrix(
        doc,
        ["Palanca", "Cómo crea valor", "Control de riesgo"],
        [
            ["Aporte base 2%", "Financia operación inicial y reserva", "Base recurrente para inscritos"],
            ["Cálculos posteriores", "Préstamos y programas especiales", "Según capacidad, topes y aprobación"],
            ["Convenios proveedores", "Más casos atendidos con el mismo fondo", "Precios negociados y cupos"],
            ["Fondo solidario", "Cubre casos de alto impacto", "Reserva mínima obligatoria"],
            ["Scoring y workflow", "Evita asignaciones improvisadas", "Comité, SLA y auditoría"],
        ],
        [1.55, 2.5, 2.3],
    )

    add_heading(doc, "6. Ejemplo financiero referencial")
    add_matrix(
        doc,
        ["Variable", "Escenario ilustrativo", "Uso directivo"],
        [
            ["Participantes", "1.500 afiliados", "Dimensionar recaudación potencial"],
            ["Aporte base inscritos", "2%", "Referencia institucional recurrente"],
            ["Base promedio ilustrativa", "300 USD equivalentes", "Permite estimar recaudación"],
            ["Aporte resultante", "6 USD equivalentes", "2% de la base ilustrativa"],
            ["Recaudación mensual", "9.000 USD equivalentes", "Techo mensual de gestión"],
            ["Reserva técnica", "20% = 1.800", "Protección contra picos de demanda"],
            ["Operación/plataforma", "10% = 900", "Costo de administrar bien"],
            ["Fondo de ayudas", "6.300", "Monto disponible antes de convenios"],
        ],
        [1.75, 2.15, 2.45],
    )
    add_paragraph(doc, "El ejemplo no cambia la política: usa 2% como aporte base y deja préstamos o programas especiales para cálculo posterior.", size=9.8, color=GRAY)

    doc.add_page_break()

    add_heading(doc, "7. Métricas para dirigir el programa")
    add_matrix(
        doc,
        ["Categoría", "Indicadores clave"],
        [
            ["Impacto social", "Personas censadas, familias alcanzadas, medicamentos, cirugías, casos resueltos, variación de calidad de vida"],
            ["Gestión operativa", "Tiempo registro-comité, tiempo comité-resolución, SLA vencidos, carga por miembro, documentos completos"],
            ["Sostenibilidad", "Aporte promedio, recaudación mensual, mora, costo por caso, descuento proveedor, reserva, siniestralidad"],
            ["Transparencia", "Casos por estado, auditoría, decisiones documentadas, métricas agregadas sin datos personales"],
        ],
        [1.7, 4.65],
    )

    add_heading(doc, "8. Hoja de ruta")
    for item in [
        "Fase 1 - Piloto controlado: validar preguntas, consentimiento, roles y primer tablero directivo.",
        "Fase 2 - Operación del comité: estados, comentarios, archivos, SLA y reportes.",
        "Fase 3 - Convenios y programas: medicamentos crónicos, cirugías, apoyo familiar y bienestar.",
        "Fase 4 - Modelo recurrente: aporte base 2%, reserva técnica, topes y medición trimestral.",
        "Fase 5 - Escala: padrón real, proveedores integrados y reportes comparativos por periodo.",
    ]:
        add_number(doc, item)

    add_heading(doc, "9. Gobernanza mínima")
    for item in [
        "Responsables con permisos definidos y trazabilidad.",
        "Política de topes por tipo de apoyo.",
        "Reserva técnica mínima antes de comprometer nuevos desembolsos.",
        "Revisión jurídica del tratamiento de datos sensibles.",
        "Reporte directivo mensual durante piloto y trimestral en operación.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "10. Cierre para directivos")
    add_paragraph(
        doc,
        "CACREF Salud permite pasar de ayuda reactiva a gestión institucional: medir, priorizar, decidir, ejecutar y auditar. Para los colaboradores PDVSA/FUTPV representa una ruta clara para presentar necesidades reales. Para CACREF representa una plataforma para construir programas recurrentes con reglas financieras, reserva, convenios y control de riesgo.",
    )
    add_callout(
        doc,
        "Mensaje final",
        "Ayudar más no debe significar gastar sin control. Con datos y gestión, CACREF puede ampliar impacto social y proteger la sostenibilidad económica de la cooperativa.",
        fill=LIGHT_BLUE,
    )
    add_heading(doc, "Decisiones inmediatas")
    for item in [
        "Aprobar piloto institucional.",
        "Nombrar comité y perfiles de acceso.",
        "Validar legalmente privacidad y consentimiento.",
        "Ratificar 2% base, topes y reserva inicial.",
        "Autorizar negociación con proveedores de salud.",
    ]:
        add_number(doc, item)


def main():
    DOCS.mkdir(exist_ok=True)
    metrics = get_metrics()
    doc = Document()
    setup_document(doc)
    add_cover(doc, metrics)
    add_pages(doc, metrics)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
